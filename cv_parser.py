"""
CV Parser - Extracts skills, experience, and keywords from CV/resume
"""

import re
from typing import List, Set, Dict
import os

# Try to import PDF parsing libraries
try:
    import PyPDF2
    PDF2_AVAILABLE = True
except ImportError:
    PDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Try to import DOCX parsing library
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class CVParser:
    def __init__(self, cv_path: str):
        """Initialize CV parser with path to CV file"""
        self.cv_path = cv_path
        self.cv_text = self._load_cv()
        self.skills = self._extract_skills()
        self.experience = self._extract_experience()
        self.keywords = self._extract_keywords()
    
    def _load_cv(self) -> str:
        """Load CV text from file (supports PDF, DOCX, and text files)"""
        if not os.path.exists(self.cv_path):
            print(f"Warning: CV file '{self.cv_path}' not found.")
            return ""
        
        # Check file type by reading header
        try:
            with open(self.cv_path, 'rb') as f:
                header = f.read(4)
                f.seek(0)
                is_pdf = header.startswith(b'%PDF')
                # Check for DOCX (ZIP archive signature with specific structure)
                is_docx = False
                if header.startswith(b'PK\x03\x04'):
                    # Might be DOCX, check if it's a valid DOCX by looking for word/document.xml
                    try:
                        import zipfile
                        with zipfile.ZipFile(self.cv_path, 'r') as zip_file:
                            file_list = zip_file.namelist()
                            is_docx = 'word/document.xml' in file_list
                    except:
                        is_docx = False
        except:
            is_pdf = False
            is_docx = False
        
        if is_pdf:
            # Try to extract text from PDF
            text = self._extract_pdf_text()
            if text:
                return text
            else:
                print("Warning: Could not extract text from PDF. Install PyPDF2 or pdfplumber.")
                return ""
        elif is_docx:
            # Try to extract text from DOCX
            text = self._extract_docx_text()
            if text:
                return text
            else:
                print("Warning: Could not extract text from DOCX. Install python-docx.")
                return ""
        else:
            # Read as text file
            try:
                with open(self.cv_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # Try with different encoding
                try:
                    with open(self.cv_path, 'r', encoding='latin-1') as f:
                        return f.read()
                except:
                    return ""
            except Exception as e:
                print(f"Error reading CV file: {e}")
                return ""
    
    def _extract_pdf_text(self) -> str:
        """Extract text from PDF file"""
        text_parts = []
        
        # Try pdfplumber first (better extraction)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(self.cv_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                return '\n'.join(text_parts)
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")
        
        # Fallback to PyPDF2
        if PDF2_AVAILABLE:
            try:
                with open(self.cv_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                return '\n'.join(text_parts)
            except Exception as e:
                print(f"PyPDF2 extraction failed: {e}")
        
        return ""
    
    def _extract_docx_text(self) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            return ""
        
        try:
            doc = Document(self.cv_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return ""
    
    def _extract_skills(self) -> Set[str]:
        """Extract technical skills from CV"""
        # Comprehensive tech skills patterns
        skills_patterns = [
            # Programming Languages
            r'\b(?:Python|Java|JavaScript|TypeScript|C\+\+|C#|Go|Rust|Ruby|PHP|Swift|Kotlin|PowerShell|Bash)\b',
            # Web Frameworks
            r'\b(?:React|Vue|Angular|Node\.js|Express|Django|Flask|Spring|Laravel)\b',
            # Cloud Platforms
            r'\b(?:AWS|Azure|GCP|Google Cloud|Hyper-V|VMware|VMware Horizon|VDI|O365|Office 365)\b',
            # Containers & Orchestration
            r'\b(?:Docker|Kubernetes|K8s|Terraform|Ansible|Infrastructure as Code|IaC)\b',
            # Databases
            r'\b(?:PostgreSQL|MySQL|MongoDB|Redis|Elasticsearch|Azure SQL|SQL Server|SSMS|SQL)\b',
            # DevOps & Tools
            r'\b(?:Git|Jenkins|CI/CD|DevOps|Agile|Scrum|SAFe|Jira|Confluence|ServiceNow|ITIL)\b',
            # Operating Systems
            r'\b(?:Linux|Unix|Windows|Windows Server|CentOS|Ubuntu|MacOS|macOS)\b',
            # AI/ML
            r'\b(?:Machine Learning|ML|AI|Data Science|TensorFlow|PyTorch)\b',
            # Infrastructure & Virtualization
            r'\b(?:Active Directory|AD|Azure AD|Exchange|SCCM|System Center|Group Policy|Nutanix|Rubrik)\b',
            # Networking
            r'\b(?:DNS|DHCP|Load Balancing|Load Balancer|VPN|Firewall|Network Monitoring|Application Gateway|FortiClient|Fortinet)\b',
            # Security
            r'\b(?:MFA|Multi-Factor Authentication|Duo|BitLocker|Identity Management|IAM|Security|Encryption)\b',
            # Monitoring & Backup
            r'\b(?:Nagios|Monitoring|Backup|Recovery|Disaster Recovery|RTO|RPO|SLA)\b',
            # Storage
            r'\b(?:SAN|Storage|NAS|Backup & Recovery)\b',
            # Communication
            r'\b(?:VOIP|Video Conferencing|Telepresence|Softphone)\b',
            # Software Tools
            r'\b(?:Autodesk|Revit|ProjectWise|Adobe Creative Suite|Adobe)\b',
            # Methodologies
            r'\b(?:ITIL|SLA Compliance|Incident Management|Service Management)\b',
        ]
        
        skills = set()
        cv_lower = self.cv_text.lower()
        
        # Search for skill patterns
        for pattern in skills_patterns:
            matches = re.finditer(pattern, self.cv_text, re.IGNORECASE)
            for match in matches:
                skills.add(match.group().lower())
        
        # Also extract from common skill sections (CORE COMPETENCIES, Skills, Technologies)
        skill_section_patterns = [
            r'(?:CORE COMPETENCIES|Skills?|Technologies?|Expertise|Competencies?)[:]\s*(.+?)(?:\n\n|\n[A-Z][A-Z]+\s|$)',
            r'(?:Cloud Platforms|Infrastructure|Security|Networking|DevOps|Databases|Service Management)[:]\s*(.+?)(?:\n[A-Z]|$)',
        ]
        
        for pattern in skill_section_patterns:
            skill_section = re.search(pattern, self.cv_text, re.IGNORECASE | re.DOTALL)
            if skill_section:
                skill_list = skill_section.group(1)
                # Split by common delimiters (comma, semicolon, bullet, dash, newline)
                individual_skills = re.split(r'[,;•\-\n\r\|]', skill_list)
                for skill in individual_skills:
                    skill = skill.strip()
                    # Filter out very short or common words
                    if skill and len(skill) > 2 and skill.lower() not in ['and', 'the', 'for', 'with', 'on', 'in', 'to']:
                        # Normalize skill name
                        skill_normalized = skill.lower().strip('.,;:()[]')
                        if skill_normalized:
                            skills.add(skill_normalized)
        
        # Extract specific technology names from anywhere in CV
        specific_techs = [
            'Active Directory', 'Azure AD', 'Azure SQL', 'VMware', 'Hyper-V', 'Nutanix',
            'Rubrik', 'Nagios', 'SCCM', 'Exchange', 'DNS', 'DHCP', 'MFA', 'VPN', 'BitLocker',
            'ServiceNow', 'Jira', 'Confluence', 'ITIL', 'O365', 'Office 365', 'FortiClient',
            'Autodesk', 'Revit', 'ProjectWise', 'Adobe', 'SLA', 'VDI', 'SAN', 'VOIP',
            'Windows Server', 'CentOS', 'Ubuntu', 'Linux', 'PowerShell', 'Bash'
        ]
        
        cv_lower = self.cv_text.lower()
        for tech in specific_techs:
            if tech.lower() in cv_lower:
                skills.add(tech.lower())
        
        return skills
    
    def _extract_experience(self) -> Dict:
        """Extract work experience information"""
        experience = {
            'years': 0,
            'companies': [],
            'roles': []
        }
        
        # Try to extract years of experience
        years_pattern = r'(\d+)[\+]?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)'
        years_match = re.search(years_pattern, self.cv_text, re.IGNORECASE)
        if years_match:
            experience['years'] = int(years_match.group(1))
        
        # Extract role titles
        role_patterns = [
            r'(?:Senior\s*)?(?:Software\s*)?(?:Engineer|Developer|Programmer)',
            r'(?:DevOps|SRE|Site Reliability) Engineer',
            r'(?:Data|ML|AI) (?:Scientist|Engineer)',
            r'(?:Product|Project|Technical) Manager',
            r'Architect',
        ]
        
        for pattern in role_patterns:
            matches = re.finditer(pattern, self.cv_text, re.IGNORECASE)
            for match in matches:
                experience['roles'].append(match.group())
        
        return experience
    
    def _extract_keywords(self) -> Set[str]:
        """Extract important keywords from CV"""
        keywords = set()
        
        # Industry-specific keywords
        industry_keywords = [
            'cloud', 'microservices', 'api', 'rest', 'graphql',
            'distributed systems', 'scalability', 'performance',
            'security', 'testing', 'automation', 'monitoring',
            'big data', 'analytics', 'full stack', 'backend', 'frontend'
        ]
        
        cv_lower = self.cv_text.lower()
        for keyword in industry_keywords:
            if keyword in cv_lower:
                keywords.add(keyword)
        
        # Add all skills as keywords
        keywords.update(self.skills)
        
        return keywords
    
    def get_skills(self) -> Set[str]:
        """Get extracted skills"""
        return self.skills
    
    def get_keywords(self) -> Set[str]:
        """Get extracted keywords"""
        return self.keywords
    
    def get_experience(self) -> Dict:
        """Get extracted experience"""
        return self.experience



